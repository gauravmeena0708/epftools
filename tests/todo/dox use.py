from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

#Page 1
logo = document.add_picture('EPFO_Logo.png', width=Inches(1.25))
last_paragraph = document.paragraphs[-1] 
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
title = document.add_heading('Regional Office, KR Puram', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p1= document.add_paragraph("""EPFO is one of the World's largest Social Security Organisations in terms of clientele \
and the volume of financial transactions undertaken. At present it maintains 24.77 crore accounts (Annual Report \
2019-20) pertaining to its members.""")
p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p2=document.add_paragraph("""The Employees' Provident Fund came into existence with the promulgation of the Employees' \
Provident Funds Ordinance on the 15th November, 1951. It was replaced by the Employees' Provident Funds Act, 1952. \
The Employees' Provident Funds Bill was introduced in the Parliament as Bill Number 15 of the year 1952 as a Bill \
to provide for the institution of provident funds for employees in factories and other establishments. """)
p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p3=document.add_paragraph("""The Act is now referred as the Employees' Provident Funds & Miscellaneous Provisions Act, \
1952 which extends to the whole of India. The Act and Schemes framed there under are administered by a tri-partite \
Board known as the Central Board of Trustees, Employees' Provident Fund,consisting of representatives of Government \
(Both Central and State), Employers, and Employees.""")
p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p4 =document.add_paragraph("""Vision: An innovation driven social security organisation aiming to extend universal \
coverage and ensuring Nirbadh (Seamless and uninterrupted) service delivery to its stakeholders through \
state-of-the-art technology.""")
p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p5 =document.add_paragraph("""Mission: To meet the evolving needs of comprehensive social security in a transparent, \
contactless, faceless and paperless manner. To ensure Nirbadh services with multi-locational and auto claim \
settlement process for disaster proofing EPFO. To ensure ease of living for members and pensioners, and \
ease of doing business for employers by leveraging Government of India’s technology platforms for reaching \
out to millions.""")
p5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


 

    
document.add_page_break()
"""
p = document.add_paragraph('An introductory paragraph about EPFO and its functions')
p.add_run(' Adding sentences. ').bold = True
p.add_run(' Bold Sentences. ')
p.add_run('Italic sentences.').italic = True
"""
#Page 2
document.add_heading('Table List', level=1)
document.add_paragraph('Pincode - Area Map', style='List Number')
document.add_paragraph('Pincode list ', style='List Number')
document.add_paragraph('DSC/E-sign Status', style='List Number')
document.add_paragraph('Establishment Modification', style='List Number')
document.add_paragraph('Industry Codewise', style='List Number')
document.add_paragraph('Industry Groupwise', style='List Number')
document.add_paragraph('Principal Employer Registeration', style='List Number')
document.add_paragraph('Voluntary 1(3) 1(4)', style='List Number')
document.add_paragraph('Exemption', style='List Number')
document.add_paragraph('Ceiling Contribution', style='List Number')
document.add_paragraph('Claims: Yearwise Workload', style='List Number')
document.add_paragraph('Claims: Yearwise F31 Workload', style='List Number')
document.add_paragraph('Claims: Yearwise F13 Workload', style='List Number')
document.add_paragraph('Claims: Yearwise Final Settlement Workload', style='List Number')
document.add_paragraph('Claims: Rejection Reasons', style='List Number')
document.add_paragraph('Claims: App E Workload', style='List Number')
document.add_paragraph('Claims: VDR Spl Workload', style='List Number')
document.add_paragraph('Claims: Basic Detail Workload', style='List Number')
document.add_paragraph('Claims: Primary Detail Workload', style='List Number')
document.add_paragraph('Claims: Other Details Detail Workload', style='List Number')
document.add_paragraph('Claims: Est Modification Detail Workload', style='List Number')
document.add_paragraph('Claims: Per Capita', style='List Number')
document.add_paragraph('Claims: Multiple Rejections', style='List Number')
document.add_paragraph('Claims: Multiple Filings', style='List Number')
document.add_paragraph('Claims: International Workers', style='List Number')
document.add_paragraph('Grievances: CPGRAMS', style='List Number')
document.add_paragraph('Grievances: EPFiGMS Yearwise', style='List Number')
document.add_paragraph('Grievances: EPFiGMS Categorywise', style='List Number')
document.add_paragraph('Compliance: 7A', style='List Number')
document.add_paragraph('Compliance: 14B', style='List Number')
document.add_paragraph('Compliance: Current Demand', style='List Number')
document.add_paragraph('Compliance: Arrear Demand(Exempted)', style='List Number')
document.add_paragraph('Compliance: Arrear Demand(Unexempted)', style='List Number')
document.add_paragraph('Compliance: CAIU', style='List Number')
document.add_paragraph('Compliance: Defaults', style='List Number')
document.add_paragraph('Compliance: Virtual Hearings', style='List Number')
document.add_paragraph('Legal: Pending court cases', style='List Number')
document.add_paragraph('Audit: Pending Audit Paras', style='List Number')
document.add_paragraph('Exemption: Exempted', style='List Number')
document.add_paragraph('Exemption: Exempted(EDLI)', style='List Number')
document.add_paragraph('Exemption: Return Filing', style='List Number')
document.add_paragraph('Exemption: Compliance Audit', style='List Number')
document.add_paragraph('Profile Completion: Aadhaar Verification', style='List Number')
document.add_paragraph('Profile Completion: Bank Verification', style='List Number')
document.add_paragraph('Profile Completion: E-Nomination', style='List Number')
document.add_paragraph('Profile Completion: Transfer of Old account to latest', style='List Number')
document.add_paragraph('Pension: Top 10 Establishment', style='List Number')
document.add_paragraph('Pension: Prayaas', style='List Number')
document.add_paragraph('Pension: Jeevan Pramaan', style='List Number')
document.add_paragraph('Pension: Pincodewise', style='List Number')
document.add_paragraph('Pension: Manual DSC Yearwise', style='List Number')
document.add_paragraph('Cash: Account 1/2/10/21/22', style='List Number')
document.add_paragraph('Admin: Strength Sanctiod vs in position designation', style='List Number')
document.add_paragraph('Admin: Strength Sanctiod vs in position Section', style='List Number')
document.add_paragraph('Admin: IT Hardware(PC,Thinclient,scanner,printer)', style='List Number')
document.add_paragraph('Admin: IT Software(OS, Office, Antivirus)', style='List Number')
document.add_paragraph('Admin: Contractual Staff', style='List Number')
document.add_paragraph('Admin: Rented Office Area', style='List Number')
document.add_paragraph('Admin: Per Capita Expenditure', style='List Number')

document.add_page_break()
pin = document.add_picture('pin.jpg', width=Inches(6.25))
last_paragraph1 = document.paragraphs[-1] 
last_paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
document.add_page_break()
records = ((1, 'LDASR', 800),(2, 'PYKRP', 700),(3,'RJRAJ',500))

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Sr'
hdr_cells[1].text = 'Office Code'
hdr_cells[2].text = 'Strength'
table.style = "Light Grid Accent 5"
for sr, code, count in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(sr)
    row_cells[1].text = code
    row_cells[2].text = str(count)
    
#document.add_page_break()

document.save('demo2.docx')

import win32com.client # install with "pip install pywin32==224"
import docx, os

wordFilename = 'demo2.docx'
fulldocpath = os.getcwd() + "\\" + wordFilename
pdfFilename = 'report.pdf'
fullpdfpath = os.getcwd() + "\\" + pdfFilename


wdFormatPDF = 17 # Word's numeric code for PDFs.
wordObj = win32com.client.Dispatch('Word.Application')
docObj = wordObj.Documents.Open(fulldocpath)
docObj.SaveAs(fullpdfpath, FileFormat=wdFormatPDF)
docObj.Close()
wordObj.Quit()
