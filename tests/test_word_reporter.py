import unittest
import os
from src.epftools.word_reporter import WordReporter
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

class TestWordReporter(unittest.TestCase):

    def setUp(self):
        self.reporter = WordReporter()
        self.output_filepath = "test_report.docx"
        self.image_path = "test_image.png"
        # Create a dummy image file for testing
        with open(self.image_path, "wb") as f:
            f.write(b"dummy_image_content")

    def tearDown(self):
        if os.path.exists(self.output_filepath):
            os.remove(self.output_filepath)
        if os.path.exists(self.image_path):
            os.remove(self.image_path)

    def test_add_heading(self):
        self.reporter.add_heading("Test Heading", level=1)
        self.assertEqual(len(self.reporter.document.paragraphs), 1)
        self.assertEqual(self.reporter.document.paragraphs[0].text, "Test Heading")

    def test_add_paragraph(self):
        self.reporter.add_paragraph("Test Paragraph")
        self.assertEqual(len(self.reporter.document.paragraphs), 1)
        self.assertEqual(self.reporter.document.paragraphs[0].text, "Test Paragraph")

    def test_add_table(self):
        data = [
            ["Header 1", "Header 2"],
            ["Row 1 Col 1", "Row 1 Col 2"],
            ["Row 2 Col 1", "Row 2 Col 2"]
        ]
        self.reporter.add_table(data)
        self.assertEqual(len(self.reporter.document.tables), 1)
        table = self.reporter.document.tables[0]
        self.assertEqual(len(table.rows), 3)
        self.assertEqual(table.cell(0, 0).text, "Header 1")
        self.assertEqual(table.cell(1, 1).text, "Row 1 Col 2")

    def test_add_picture(self):
        self.reporter.add_picture(self.image_path, width=Inches(1))
        # Assert that a run containing a picture is added
        self.assertTrue(any(run.element.xml.find('{http://schemas.openxmlformats.org/drawingml/2006/picture}pic') is not None
                            for p in self.reporter.document.paragraphs for run in p.runs))

    def test_add_page_break(self):
        self.reporter.add_page_break()
        # Check if a page break element is added (this is a bit indirect)
        self.assertTrue(any(p.runs and p.runs[-1].text == '' and p.runs[-1].element.xml.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br') is not None
                            for p in self.reporter.document.paragraphs))

    def test_save(self):
        self.reporter.add_paragraph("Content to save")
        self.reporter.save(self.output_filepath)
        self.assertTrue(os.path.exists(self.output_filepath))

if __name__ == '__main__':
    unittest.main()
