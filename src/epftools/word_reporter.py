from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

logger = logging.getLogger(__name__)

class WordReporter:
    """
    A class for generating Word documents.
    """

    def __init__(self):
        self.document = Document()

    def add_heading(self, text, level=1, alignment=WD_ALIGN_PARAGRAPH.LEFT):
        """
        Adds a heading to the document.
        """
        heading = self.document.add_heading(text, level)
        heading.alignment = alignment
        logger.info(f"Added heading: {text}")

    def add_paragraph(self, text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
        """
        Adds a paragraph to the document.
        """
        paragraph = self.document.add_paragraph(text)
        paragraph.alignment = alignment
        logger.info(f"Added paragraph: {text}")

    def add_table(self, data, style="Light Grid Accent 5"):
        """
        Adds a table to the document.
        `data` should be a list of lists, where the first list is the header.
        """
        if not data or not isinstance(data, list) or not all(isinstance(row, list) for row in data):
            logger.warning("Invalid data format for table. Skipping table creation.")
            return

        table = self.document.add_table(rows=1, cols=len(data[0]))
        table.style = style

        # Add header row
        hdr_cells = table.rows[0].cells
        for i, header_text in enumerate(data[0]):
            hdr_cells[i].text = str(header_text)

        # Add data rows
        for row_data in data[1:]:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        logger.info(f"Added table with {len(data) - 1} rows and {len(data[0])} columns.")

    def add_picture(self, image_path, width=None, height=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
        """
        Adds a picture to the document.
        """
        try:
            picture = self.document.add_picture(image_path, width=width, height=height)
            last_paragraph = self.document.paragraphs[-1]
            last_paragraph.alignment = alignment
            logger.info(f"Added picture: {image_path}")
        except Exception as e:
            logger.error(f"Failed to add picture {image_path}: {e}")

    def add_page_break(self):
        """
        Adds a page break to the document.
        """
        self.document.add_page_break()
        logger.info("Added page break.")

    def save(self, filepath):
        """
        Saves the document to the specified filepath.
        """
        try:
            self.document.save(filepath)
            logger.info(f"Word document saved to {filepath}")
        except Exception as e:
            logger.error(f"Failed to save Word document to {filepath}: {e}")